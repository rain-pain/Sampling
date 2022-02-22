from math import fabs, sqrt, log, exp
import func as fc
from docx import Document
import os

def main():
    doc = Document()
    data, file_name = fc.file_data()
    if not os.path.isdir(f'{file_name[:-4]}'):
        os.mkdir(f'{file_name[:-4]}')

    n = int(len(data) / 2)
    head = doc.add_heading('Двумерная выборка', 1)
    head.alignment = 1

    x, y = [], []
    for i in range(len(data)):
        if i % 2 == 0:
            x.append(data[i])
        else:
            y.append(data[i])   

    p = doc.add_paragraph('\n       Исходные данные: ')
    for i in range(n - 1):
        p.add_run(f'({round(x[i], 2)}; {round(y[i], 2)}), ')
    p.add_run(f'({round(x[n - 1], 2)}; {round(y[n - 1], 2)})\n')

    sred_x, sred_y, nach_mom_x, nach_mom_y, nach_mom = 0, 0, 0, 0, 0

    table = doc.add_table(rows=n + 1, cols=6)
    table.style = 'Table Grid'
    c = table.rows[0].cells
    c[0].text = '№'
    c[1].text = 'x'
    c[2].text = 'y'
    c[3].text = 'x²'
    c[4].text = 'y²'
    c[5].text = 'x*y'

    for i in range(n):
        sred_x += x[i]
        sred_y += y[i]
        nach_mom_x += x[i] ** 2
        nach_mom_y += y[i] ** 2
        nach_mom += x[i] * y[i]

        c = table.rows[i + 1].cells
        c[0].text = str(i + 1)
        c[1].text = str(round(x[i], 4))
        c[2].text = str(round(y[i], 4))
        c[3].text = str(round(x[i] ** 2, 3))
        c[4].text = str(round(y[i] ** 2, 3))
        c[5].text = str(round(x[i] * y[i], 3))

    math_ojid_x = sred_x = round(sred_x / n, 2)
    math_ojid_y = sred_y = round(sred_y / n, 2)
    doc.add_paragraph('\n       Для решения задачи удобно воспользоваться приведенной\
 выше таблицей. Значения в 3-ем, 4-ом и 5-ом столбцах вычисляются по формулам, \
приведенными в первой строке таблицы. Таким\
 образом получены:\n\n - оценки математических ожиданий по каждой переменной')
    formula = r'"m(x) = x' + r'$^{cp}=\frac{1}{n}  %s$' + f' * ∑xᵢ"'
    fc.formulas(formula, math_ojid_x, 'math_ojid_x', file_name, ',  i = 1...n')
    doc.add_picture(f'./{file_name[:-4]}/additional/math_ojid_x.png')
    formula = r'"m(y) = y' + r'$^{cp}=\frac{1}{n}  %s$' + f' * ∑yᵢ"'
    fc.formulas(formula, math_ojid_y, 'math_ojid_y', file_name, ',  i = 1...n')
    doc.add_picture(f'./{file_name[:-4]}/additional/math_ojid_y.png')

    nach_mom_x = round(nach_mom_x / n, 5)
    nach_mom_y = round(nach_mom_y / n, 5)
    doc.add_paragraph(f' - оценки начальных моментов второго порядка по каждой переменной: ')
    formula = '"α(x)' + r'$=\frac{1}{n}  %s$' + ' * ∑xᵢ²"'
    fc.formulas(formula, nach_mom_x, 'nach_mom_x', file_name, ',  i = 1...n')
    doc.add_picture(f'./{file_name[:-4]}/additional/nach_mom_x.png')
    formula = '"α(y)' + r'$=\frac{1}{n}  %s$' + ' * ∑yᵢ²"'
    fc.formulas(formula, nach_mom_y, 'nach_mom_y', file_name, ',  i = 1...n')
    doc.add_picture(f'./{file_name[:-4]}/additional/nach_mom_y.png')

    nach_mom = round(nach_mom / n, 5)
    doc.add_paragraph(f' - оценка смешанного начального момента второго порядка: ')
    formula = '"α(x,y)' + r'$=\frac{1}{n}  %s$' + ' * ∑yᵢxᵢ"'
    fc.formulas(formula, nach_mom, 'nach_mom', file_name, ',  i = 1...n')
    doc.add_picture(f'./{file_name[:-4]}/additional/nach_mom.png')

    disp_x = round(n / (n - 1) * nach_mom_x - n / (n - 1) * (math_ojid_x ** 2), 5)
    disp_y = round(n / (n - 1) * nach_mom_y - n / (n - 1) * (math_ojid_y ** 2), 5)
    doc.add_paragraph(f'\n      На основе этих данных легко вычислить оценки дисперсий:')
    formula = '"D(x) = S₀²(x) =' + r'$\frac{1}{n-1} %s$' + f'∑xᵢ² - ' + r'$\frac{n}{n-1} %s$' + f'x' + r'$^{ср²}$"'
    fc.formulas(formula, disp_x, 'dispersion_x', file_name, ',  i = 1...n')
    doc.add_picture(f'./{file_name[:-4]}/additional/dispersion_x.png')
    doc.paragraphs[-1].alignment = 1
    formula = '"D(y) = S₀²(y) =' + r'$\frac{1}{n-1} %s$' + f'∑yᵢ² - ' + r'$\frac{n}{n-1} %s$' + f'y' + r'$^{ср²}$"'
    fc.formulas(formula, disp_y, 'dispersion_y', file_name, ',  i = 1...n')
    doc.add_picture(f'./{file_name[:-4]}/additional/dispersion_y.png')
    doc.paragraphs[-1].alignment = 1

    korelyac = round(n / (n - 1) * nach_mom - n / (n - 1) * (math_ojid_x * math_ojid_y), 6)
    doc.add_paragraph(f'и оценку корреляционного момента: ')
    formula = '"K(x,y) = ' + r'$\frac{1}{n-1} %s$' + f'∑xᵢyᵢ - ' + r'$\frac{n}{n-1} %s$' + \
                  f'x' + r'$^{ср}$' +f'y' + r'$^{ср}$"'
    fc.formulas(formula, korelyac, 'korelyac', file_name, ',  i = 1...n')
    doc.add_picture(f'./{file_name[:-4]}/additional/korelyac.png')
    doc.paragraphs[-1].alignment = 1

    toch_ocenka = round(korelyac / sqrt(disp_x * disp_y), 3)
    doc.add_paragraph(f'        Вычислим точечную оценку коэффициент корреляции по формуле:')
    formula = r'"R(x,y) = ' + r'$\frac{K(x,y)}{\sqrt{S₀²(x)S₀²(y)}}  %s$"'
    fc.formulas(formula, toch_ocenka, 'toch_ocenka', file_name)
    doc.add_picture(f'./{file_name[:-4]}/additional/toch_ocenka.png')
    doc.paragraphs[-1].alignment = 1

    doc.add_paragraph(f'\n      Вычислим интервальную оценку коэффициента корреляции с \
надежностью γ = 0,95. Для этого в таблице функции Лапласа \
найдем значение, равное γ/2 = 0.475. Далее определим соответствующее \
аему занчение аргумента: zᵧ = argФ(0.475) = 1.96. \n      Вычислим вспомогательные значения a и b: ')
    a = round(0.5 * log((1 + toch_ocenka) / (1 - toch_ocenka)) - 1.96 / sqrt(n - 3), 4)
    b = round(0.5 * log((1 + toch_ocenka) / (1 - toch_ocenka)) + 1.96 / sqrt(n - 3), 4)
    formula = r'"a = 0.5ln(' + r'$\frac{1+R(x,y)}{1-R(x,y)})-\frac{zᵧ}{\sqrt{n-3}}  %s$"'
    fc.formulas(formula, a, 'a', file_name)
    doc.add_picture(f'./{file_name[:-4]}/additional/a.png')
    formula = r'"b = 0.5ln(' + r'$\frac{1+R(x,y)}{1-R(x,y)})+\frac{zᵧ}{\sqrt{n-3}}  %s$"'
    fc.formulas(formula, b, 'b', file_name)
    doc.add_picture(f'./{file_name[:-4]}/additional/b.png')

    i_a = round((exp(2*a) - 1)/(exp(2*a) + 1), 3)
    i_b = round((exp(2*b) - 1)/(exp(2*b) + 1), 3)
    doc.add_paragraph('\n       Таким образом, доверительный интервал для коэффициента корреляции имеет вид: ')
    formula = r'"I(R) = [' + r'$\frac{e^{2a}-1}{e^{2a}+1}}  %s$' + ';' + r'$\frac{e^{2b}-1}{e^{2b}+1}}  %s$' + ']"'
    fc.formulas(formula, f'[{i_a};{i_b}]', 'I', file_name)
    doc.add_picture(f'./{file_name[:-4]}/additional/I.png')
    doc.paragraphs[-1].alignment = 1

    doc.add_paragraph('         Проверим гипотезу об отсутствии корреляционной зависимости: ')
    doc.add_paragraph('H₀ : R(x,y) = 0;\nH₁ : R(x,y) ≠ 0')
    doc.paragraphs[-1].alignment = 1

    z = round((fabs(toch_ocenka) * sqrt(n)) / (1 - toch_ocenka ** 2), 4)
    doc.add_paragraph(f'        Так как объем выборки велик (n ≥ 50 ), то вычислим значение критерия по формуле: ')
    formula = r'"Z = ' + r'$\frac{|R(x,y))| * \sqrt{n}}{1 - (R(x,y))^{2}}  %s$"'
    fc.formulas(formula, z, 'z', file_name)
    doc.add_picture(f'./{file_name[:-4]}/additional/z.png')
    doc.paragraphs[-1].alignment = 1

    doc.add_paragraph('         Определим значение Zₐ из таблицы функции Лапласа: ')
    formula = r'"Zₐ = argФ(' + r'$\frac{1-α}{2}  %s$' + ')"'
    fc.formulas(formula, 1.96, 'z2', file_name)
    doc.add_picture(f'./{file_name[:-4]}/additional/z2.png')
    doc.paragraphs[-1].alignment = 1

    if z < 1.96:
        doc.add_paragraph('         Так как Z < Zₐ, то гипотеза H₀ принимается, т.е. величины X и Y некоррелированны.')
    else:
        doc.add_paragraph('         Так как Z > Zₐ, то гипотеза H₀ не принимается, т.е. величины X и Y коррелированны.')

    a1 = round(korelyac / disp_x, 3)
    a0 = round(sred_y - a1 * sred_x, 3)
    doc.add_paragraph(f'        Вычислим оценки параметров a₁ и a₀ линии регрессии y(x) = a₀ + a₁*x по формуле:')
    formula = r'"a₁ = ' + r'$\frac{K(x,y)}{S₀²(x)}  %s$"'
    fc.formulas(formula, a1, 'a1', file_name)
    doc.add_picture(f'./{file_name[:-4]}/additional/a1.png')
    formula = r'"a₀ = ' + r'$y^{ср}-a₁x^{ср}   %s$"'
    fc.formulas(formula, a0, 'a0', file_name)
    doc.add_picture(f'./{file_name[:-4]}/additional/a0.png')

    doc.add_paragraph(f'        Уравнение линии регрессии имеет вид :')
    if a1 >= 0:
        p = doc.add_paragraph(f'y(x) = {a0} + {a1}*x')
        p.alignment = 1
    else:
        p = doc.add_paragraph(f'y(x) = {a0}{a1}*x')
        p.alignment = 1

    doc.add_paragraph('        Построим диаграмму рассеивания, изобразив значения исходной \
двумерной выборки в виде точек с координатами ( xᵢ, yᵢ ) на плоскости \
в декартовой системе координат, и линию регрессии. ')
    x_graph = [min(x) - 1, max(x) + 1]
    y_graph = [a0 + a1 * x_graph[0], a0 + a1 * x_graph[1]]
    fc.graph_save(x, y, x_graph, y_graph, file_name)

    doc.add_picture(f'./{file_name[:-4]}/Image_{file_name[:-4]}.png')
    doc.save(f'./{file_name[:-4]}/{file_name[:-4]}.docx')
    
if __name__ == '__main__':
    main()
